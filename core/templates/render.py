"""DraftOrdinance / AmendmentDiff → .docx 출력.

한글 오피스가 .docx를 그대로 읽을 수 있어서, 임희연님은 출력 .docx를 한글에서 열어
'다른 이름으로 저장 → .hwpx'로 한 번만 변환하면 의회 양식과 호환된다.

설계 원칙:
- 한국 입법 양식 (제안이유 → 주요내용 → 조문/신·구조문대비표 → 부칙) 그대로
- 시스템 기본 한글 폰트(맑은 고딕)로 안전하게. 의회 표준 폰트가 따로 있다면
  추후 setting으로 빼면 된다.
- 표는 신·구조문대비표용으로만 사용 (3열: 현행 / 개정안 / 비고)
- 위험한 자동 줄바꿈·페이지 분리 없음 — 사용자가 한글에서 마무리 손볼 여지.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from core.models import AmendmentDiff, DraftOrdinance


KOREAN_FONT = "맑은 고딕"  # Windows 기본 — 환경에 항상 있음


def render_new_ordinance(
    draft: DraftOrdinance,
    output_path: str | Path,
) -> Path:
    """제정조례안을 .docx로 출력. 셀프 점검 결과는 *별도 파일*로 분리되어 첨부 안 함."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _apply_default_font(doc)
    _apply_page_margins(doc)

    _add_title(doc, draft.title)
    if draft.delegation_law:
        _add_meta_line(doc, f"위임 근거: {draft.delegation_law}")

    _add_section_header(doc, "Ⅰ. 제안이유")
    _add_paragraph_block(doc, draft.proposal_reason)

    _add_section_header(doc, "Ⅱ. 주요내용")
    for line in draft.main_contents:
        _add_indented_item(doc, line)

    _add_section_header(doc, "Ⅲ. 조문")
    for art in draft.articles:
        label = (art.get("label") or "").strip()
        title = (art.get("title") or "").strip()
        body = (art.get("body") or "").strip()
        head = f"{label}{title}" if title.startswith("(") else f"{label}({title})" if title else label
        _add_article(doc, head, body)

    _add_section_header(doc, "부칙")
    _add_paragraph_block(doc, draft.addendum)

    doc.save(output_path)
    return output_path


def render_amendment(
    draft: DraftOrdinance,
    diffs: list[AmendmentDiff],
    output_path: str | Path,
    *,
    proposer_info: str = "",
    bill_no: str = "",
    propose_date: str = "",
    proposer_list: str = "",
    related_laws: list[dict] | None = None,
    directives: list[str] | None = None,
) -> Path:
    """일부개정조례안 .docx — 경기도의회 표준 양식.

    구조:
      1) 표지: 조례명 + (○○ 의원 대표발의) + 의안번호 표
      2) 본문: 1.제안이유 / 2.주요내용 / 3·4·5 덧붙임 안내
      3) 개정조례안 본문: 개정 지시문 + 부칙
      4) 신·구조문 대비표: 2열 표 (현행 / 개정안)
      5) 관계법령 발췌서: □ 법령명 + 본문

    Args:
        proposer_info: "(○○ 의원 대표발의)" — 표지 부제. 비우면 표시 안 함.
        bill_no, propose_date, proposer_list: 의안번호 표 칸들. 비워두면 "(미정)" 표시.
        related_laws: [{"name": "...", "body": "..."}, ...] 관계법령 발췌. 비우면 빈 양식.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _apply_default_font(doc)
    _apply_page_margins(doc)

    # ============ 1) 표지 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run(draft.title)
    run.bold = True
    run.font.size = Pt(20)

    if proposer_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(proposer_info)
        run.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    _add_cover_bill_table(
        doc,
        bill_no=bill_no or "○○○○",
        propose_date=propose_date or "2026년 ○월 ○일",
        proposer_list=proposer_list or "○○○ 의원(○명)",
    )

    doc.add_page_break()

    # ============ 2) 본문 (제안이유 / 주요내용 / 덧붙임 안내) ============
    _add_amendment_numbered_header(doc, "1. 제안이유")
    _add_paragraph_block(doc, draft.proposal_reason)

    doc.add_paragraph()

    _add_amendment_numbered_header(doc, "2. 주요내용")
    for line in draft.main_contents:
        _add_indented_item(doc, line)

    doc.add_paragraph()

    _add_amendment_numbered_header(doc, "3. 개정조례안: 덧붙임")
    _add_amendment_numbered_header(doc, "4. 신·구조문 대비표: 덧붙임")
    _add_amendment_numbered_header(doc, "5. 관계법령 발췌서: 덧붙임")

    doc.add_page_break()

    # ============ 3) 개정조례안 본문 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("경기도 조례 제      호")
    run.font.size = Pt(11)
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(draft.title)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph()

    base_title = draft.title.replace(" 일부개정조례안", "").strip()
    p = doc.add_paragraph()
    run = p.add_run(f"{base_title}의 일부를 다음과 같이 개정한다.")
    run.font.size = Pt(11)
    doc.add_paragraph()

    # 개정 지시문 — 호출자가 LLM으로 생성해서 전달한 표준 지시문 사용 (옛 휴리스틱 fallback 유지).
    directive_lines = directives if directives is not None else _build_amendment_directives(diffs)
    for directive in directive_lines:
        p = doc.add_paragraph(directive)
        p.paragraph_format.space_after = Pt(4)

    # 부칙
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("부      칙")
    run.bold = True
    run.font.size = Pt(13)
    # 중복 제거: LLM addendum 본문 안에 "부칙"/"부 칙" 표시가 들어있으면 줄을 떼어냄.
    addendum_body = (draft.addendum or "이 조례는 공포한 날부터 시행한다.").strip()
    addendum_lines = [l.strip() for l in addendum_body.splitlines() if l.strip()]
    addendum_lines = [
        l for l in addendum_lines
        if l not in ("부칙", "부 칙", "부      칙", "부  칙")
    ]
    _add_paragraph_block(doc, "\n".join(addendum_lines))

    doc.add_page_break()

    # ============ 4) 신·구조문 대비표 (2열) ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("신·구조문 대비표")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    if diffs:
        _add_2col_diff_table(doc, diffs)
    else:
        _add_paragraph_block(doc, "(변경 사항이 식별되지 않았습니다.)")

    doc.add_page_break()

    # ============ 5) 관계법령 발췌서 ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("관계법령 발췌서")
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph()

    if related_laws:
        for law in related_laws:
            name = (law.get("name") or "").strip()
            body = (law.get("body") or "").strip()
            if name:
                p = doc.add_paragraph()
                run = p.add_run(f"□ {name}")
                run.bold = True
                run.font.size = Pt(12)
            if body:
                _add_paragraph_block(doc, body)
            doc.add_paragraph()
    else:
        _add_paragraph_block(
            doc,
            "(관련 상위법령의 인용 조항을 첨부하세요 — 예: 「지방자치법」 제28조, "
            "「노인복지법」 제4조 등.)",
        )

    doc.save(output_path)
    return output_path


# ---------------- low-level helpers ----------------


def _apply_default_font(doc: Document) -> None:
    """문서 기본 폰트를 한글 친화 폰트로 지정.

    python-docx는 Western/East Asian 폰트를 별도로 다룬다. 한글 글자는 East Asian
    슬롯에 들어가므로 rFonts의 eastAsia 속성을 같이 설정해야 모든 글자가 통일된 폰트로
    렌더된다.
    """
    normal = doc.styles["Normal"]
    normal.font.name = KOREAN_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)


def _apply_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _add_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)


def _add_meta_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


def _add_section_header(doc: Document, text: str) -> None:
    doc.add_paragraph()  # 빈 줄
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _add_paragraph_block(doc: Document, text: str) -> None:
    """줄바꿈 단위로 문단 분리해 추가."""
    if not text:
        return
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        doc.add_paragraph(line)


def _add_indented_item(doc: Document, text: str) -> None:
    """주요내용 등의 항목: 들여쓰기 + 줄바꿈."""
    p = doc.add_paragraph(text.strip())
    p.paragraph_format.left_indent = Cm(0.6)


def _add_article(doc: Document, head: str, body: str) -> None:
    """조문 한 개를 머리 + 본문으로 출력. 본문은 항(①), 호(1.) 단위 줄바꿈 유지."""
    p = doc.add_paragraph()
    run = p.add_run(head)
    run.bold = True
    if body:
        # 본문 추가 — head 다음 줄부터
        for line in body.split("\n"):
            line = line.rstrip()
            if not line:
                continue
            doc.add_paragraph(line).paragraph_format.left_indent = Cm(0.5)
    doc.add_paragraph()  # 조 사이 빈 줄


def _add_cover_bill_table(doc, *, bill_no: str, propose_date: str, proposer_list: str) -> None:
    """표지의 의안번호 표 (의안번호 / 발의연월일 / 발의자)."""
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"

    # 1행 1열: "의 안 번 호" + 의안번호 (가운데 정렬)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("의   안\n번   호")
    run.bold = True
    run.font.size = Pt(11)

    cell = table.cell(0, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(bill_no)
    run.bold = True
    run.font.size = Pt(14)

    # 1행 3-4열: 발의연월일
    cell = table.cell(0, 2)
    p = cell.paragraphs[0]
    run = p.add_run("발의연월일 :")
    run.font.size = Pt(11)

    cell = table.cell(0, 3)
    p = cell.paragraphs[0]
    run = p.add_run(propose_date)
    run.font.size = Pt(11)

    # 2행 3-4열: 발의자
    cell = table.cell(1, 2)
    p = cell.paragraphs[0]
    run = p.add_run("발  의  자 :")
    run.font.size = Pt(11)

    cell = table.cell(1, 3)
    _format_table_cell(cell, proposer_list)

    # 1열은 두 줄을 차지 — w:vMerge
    _vmerge_cells(table, col=0, rows=[0, 1])
    _vmerge_cells(table, col=1, rows=[0, 1])


def _vmerge_cells(table, *, col: int, rows: list[int]) -> None:
    """지정 셀들을 세로 병합 (vMerge)."""
    from docx.oxml import OxmlElement
    for i, r in enumerate(rows):
        cell = table.cell(r, col)
        tcPr = cell._tc.get_or_add_tcPr()
        vMerge = OxmlElement("w:vMerge")
        if i == 0:
            vMerge.set(qn("w:val"), "restart")
        tcPr.append(vMerge)


def _add_amendment_numbered_header(doc, text: str) -> None:
    """일부개정조례안 본문 1./2./... 헤더."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _build_amendment_directives(diffs: list[AmendmentDiff]) -> list[str]:
    """diffs → 개정조례안 본문에 들어갈 평문 개정 지시문 목록.

    실제 자치법규 양식의 패턴을 단순 휴리스틱으로 재현:
      - modified : "제N조를 다음과 같이 한다. [개정안 본문]"
      - new      : "제N조를 다음과 같이 신설한다. [신설 본문]"
      - deleted  : "제N조를 삭제한다."

    더 세밀한 "중 'AAA'를 'BBB'로 한다" 패턴은 LLM 후속 개선 여지.
    """
    out: list[str] = []
    for d in diffs:
        label = (d.article_label or "").strip()
        ct = d.change_type
        rev = (d.revised_text or "").strip()
        cur = (d.current_text or "").strip()
        if ct == "deleted":
            out.append(f"{label}을 삭제한다.")
        elif ct == "new":
            if rev:
                out.append(f"{label}을 다음과 같이 신설한다.")
                out.append(rev)
            else:
                out.append(f"{label}을 신설한다.")
        else:  # modified
            if rev:
                out.append(f"{label}을 다음과 같이 한다.")
                out.append(rev)
            elif cur:
                out.append(f"{label}을 다음과 같이 한다.")
                out.append(cur)
            else:
                out.append(f"{label}을 개정한다.")
    return out


def _add_2col_diff_table(doc, diffs: list[AmendmentDiff]) -> None:
    """신·구조문 대비표 — 2열 (현행 / 개정안). *임병수 「법률입안상식」 표준*.

    규칙:
      - 신설 행: 현행란 *빈칸*, 개정안란 본문 전체 + 밑줄.
      - 삭제 행: 현행란 본문 전체 + 밑줄, 개정안란 *빈칸*.
      - 수정 행: 양쪽 모두 *조 전체* 표시. LLM이 `__...__` 마커로 감싼 부분만 밑줄.
      - 변경 안 된 부분은 말줄임표(`…………`)로 (대시 X).
      - 라벨은 *조 단위*로만 (제2조, 제8조의2). 항·호는 셀 안 줄바꿈으로.

    표 사양:
      - 두 열의 너비를 *균등 분할* (각 8cm).
      - autofit = False (셀 폭 고정).
      - 셀 자동 페이지 분할 허용 (`cantSplit = False`).
    """
    # 안내 박스 — 실무에서는 법령안 편집기로 마무리 작업해야 함.
    info = doc.add_paragraph()
    info_run = info.add_run(
        "※ 본 신·구조문대비표는 LLM이 생성한 *초안*입니다. "
        "정식 의안 제출 시에는 법령안 편집기(법제처·국회법제실 등)에서 "
        "현행 조문 본문을 정확히 가져와 *재작성*하시기를 권장합니다."
    )
    info_run.italic = True
    info_run.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ("현   행", "개정안")):
        cell.width = Cm(8.0)
        _format_table_cell(cell, label, bold=True, center=True)

    for d in diffs:
        row = table.add_row().cells
        # 행마다 셀 너비 재명시 (워드는 row마다 width 잃을 수 있음)
        row[0].width = Cm(8.0)
        row[1].width = Cm(8.0)

        if d.change_type == "new":
            # 좌측 마커 <신   설> 자체에도 밑줄 (희연 팀장 지시)
            current = "__<신   설>__"
            rev_body = (d.revised_text or "").strip()
            # 본문 전체 underline — 마커가 이미 있으면 추가 안 함
            if rev_body:
                if not rev_body.startswith("__") or not rev_body.endswith("__"):
                    revised = f"__{rev_body}__"
                else:
                    revised = rev_body
            else:
                revised = ""
        elif d.change_type == "deleted":
            cur_body = (d.current_text or "").strip()
            if cur_body:
                if not cur_body.startswith("__") or not cur_body.endswith("__"):
                    current = f"__{cur_body}__"
                else:
                    current = cur_body
            else:
                current = ""
            # 우측 마커 <삭   제>에도 밑줄
            revised = "__<삭   제>__"
        else:
            # modified — 양쪽 모두 *조 전체*, LLM이 만든 마커 그대로
            current = (d.current_text or "").strip()
            revised = (d.revised_text or "").strip()
            # 자동 보강: 조 번호만 바뀐 케이스 → 좌측 옛 라벨 + 우측 새 라벨 둘 다 밑줄
            import re as _re
            pat = _re.compile(r"제\d+조(?:의\d+)?")

            def _first_label(s: str) -> str:
                # 빈 줄·기존 마커·앞 공백을 무시하고 첫 조 라벨 찾기.
                for line in s.split("\n"):
                    cleaned = _re.sub(r"^[\s_]+", "", line)
                    m = pat.match(cleaned)
                    if m:
                        return m.group(0)
                return ""

            old_label = _first_label(current)
            new_label = _first_label(revised)
            if old_label and new_label and old_label != new_label:
                if f"__{old_label}__" not in current:
                    current = current.replace(old_label, f"__{old_label}__", 1)
                if f"__{new_label}__" not in revised:
                    revised = revised.replace(new_label, f"__{new_label}__", 1)

        # 변경 없음 요약(`①∼③ (생 략)` 등)을 조 헤더 줄과 띄어쓰기로 연결
        current = _compact_unchanged_summary(current)
        revised = _compact_unchanged_summary(revised)

        _format_table_cell_with_underline(row[0], current)
        _format_table_cell_with_underline(row[1], revised)


def _has_marker(text: str) -> bool:
    """텍스트에 `__...__` 마커가 이미 있는지 — 있으면 자동 감싸기 스킵."""
    return "__" in text


def _compact_unchanged_summary(text: str) -> str:
    """*변경 없음* 요약(`①∼③ (생 략)`, `(현행과 같음)` 등)을 조 헤더 줄과 띄어쓰기로 연결.

    실무 표준: `제3조(도지사의 책무) ①∼③ (생 략)` 처럼 한 줄로 표시.
    """
    import re

    lines = text.split("\n")
    out: list[str] = []
    # 짧은 요약 패턴 — `①∼② (생 략)`, `① ∼ ③ (현행과 같음)`, `(생 략)` 등
    summary_pat = re.compile(
        r"^[①②③④⑤⑥⑦⑧⑨⑩\d.~∼\s]{0,15}\(\s*(?:생\s*략|현행과\s*같음)\s*\)\s*$"
    )
    for line in lines:
        stripped = line.strip()
        if out and stripped and summary_pat.match(stripped):
            # 앞 줄과 띄어쓰기로 연결
            out[-1] = out[-1].rstrip() + " " + stripped
        else:
            out.append(line)
    return "\n".join(out)


def _format_table_cell_with_underline(cell, text: str) -> None:
    """셀에 텍스트를 채우되 `__...__` 마커로 감싼 부분에 underline=True 적용.

    예: "법 제75조의5__제5항__-----."  →  세 run으로:
      run("법 제75조의5"), run("제5항", underline=True), run("-----.")

    멀티라인 마커(`__...첫 줄\n둘째 줄...__`)도 지원 — 줄바꿈을 포함한
    *전체 마커 구간*을 찾아 underline 처리한 뒤 각 줄을 별도 paragraph로.
    """
    import re

    cell.text = ""  # 기본 빈 paragraph 제거
    p = cell.paragraphs[0]
    first_para = True

    # 전체 텍스트에서 __...__ 마커를 찾아 (text, is_underline) 세그먼트 리스트로 분해.
    # 비탐욕 + DOTALL — 마커 안에 \n 있어도 매칭.
    segments: list[tuple[str, bool]] = []
    last = 0
    for m in re.finditer(r"__(.+?)__", text, flags=re.DOTALL):
        if m.start() > last:
            segments.append((text[last : m.start()], False))
        segments.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        segments.append((text[last:], False))

    # 세그먼트를 줄바꿈 기준으로 다시 split하면서 paragraph 분할.
    for seg_text, is_under in segments:
        # \n으로 split하면 첫 번째는 현재 단락에 이어지고, 이후 항목은 새 단락.
        lines = seg_text.split("\n")
        for li, line_text in enumerate(lines):
            if li > 0:
                # 새 단락 만들기
                p = cell.add_paragraph()
                first_para = False
            if not line_text:
                continue
            run = p.add_run(line_text)
            run.font.size = Pt(10)
            if is_under:
                run.underline = True


def _add_diff_table(doc: Document, diffs: list[AmendmentDiff]) -> None:
    """신·구조문대비표: 3열 표 (현행 / 개정안 / 비고).

    한국 입법 양식에서 비고에는 조항 라벨과 변경 사유를 함께 표기.
    """
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for cell, label in zip(hdr, ("현 행", "개정안", "비 고")):
        _format_table_cell(cell, label, bold=True, center=True)

    for d in diffs:
        row = table.add_row().cells
        current = d.current_text or "<신설>"
        revised = d.revised_text or "<삭제>"
        note_lines = [d.article_label]
        if d.note:
            note_lines.append(d.note)
        if d.change_type and d.change_type != "modified":
            note_lines.append(f"[{d.change_type}]")

        _format_table_cell(row[0], current)
        _format_table_cell(row[1], revised)
        _format_table_cell(row[2], "\n".join(note_lines))


def _format_table_cell(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = ""  # 기본 빈 문단 제거 후 재구성
    p = cell.paragraphs[0]
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p = cell.add_paragraph()
            if center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = bold
        run.font.size = Pt(10)


# ---------------- self-check appendix ----------------
#
# 작성자(정책지원관 페르소나)의 셀프 점검 결과를 초안 .docx의 마지막 페이지에
# *참고용*으로 첨부. 본격 검토 트랙의 검토보고서와 다른 가벼운 sanity check.


def _add_self_check_section(doc, self_check) -> None:
    """SelfCheckResult를 마지막 페이지에 첨부."""
    from docx.shared import RGBColor

    # 헤더
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[참고] 자동 셀프 점검 결과")
    run.bold = True
    run.font.size = Pt(15)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("※ 본격 검토가 아닌 작성자 주의 환기용 sanity check")
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    if self_check.overall_note:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"⚠ 주의 환기 : {self_check.overall_note}")
        run.bold = True
        run.font.size = Pt(12)

    # 1. 결정값 반영
    _sc_h(doc, "1. 결정값 반영 점검")
    refl = self_check.decision_reflection
    if refl.covered:
        _sc_item(doc, "○ 반영 확인 : " + ", ".join(refl.covered))
    if refl.issues:
        _sc_item(doc, "○ ⚠ 어긋남/누락된 결정값")
        for issue in refl.issues:
            ref = f" ({issue.article_ref})" if issue.article_ref else ""
            _sc_sub(doc, f"- [{issue.decision}]{ref} {issue.concern}")
    if not refl.issues:
        _sc_item(doc, "○ ✅ 누락·어긋남 없음")

    # 2. 상위법 저촉
    _sc_h(doc, "2. 상위법 명백한 저촉 (sanity check)")
    if self_check.higher_law_conflicts:
        for c in self_check.higher_law_conflicts:
            _sc_item(doc, f"○ ⚠ {c.article} — {c.concern}  (신뢰도: {c.confidence})")
    else:
        _sc_item(doc, "○ ✅ 명백한 충돌 미확인")

    # 3. 환각 의심
    _sc_h(doc, "3. 인용된 법률·조례 환각 의심")
    if self_check.citation_hallucination_suspects:
        for s in self_check.citation_hallucination_suspects:
            _sc_item(doc, f"○ ⚠ '{s.citation_text}' — {s.concern}")
    else:
        _sc_item(doc, "○ ✅ 환각 의심 인용 없음")

    # 4. 표준 구조 누락
    _sc_h(doc, "4. 표준 구조 누락 점검")
    if self_check.missing_standard_components:
        for m in self_check.missing_standard_components:
            _sc_item(doc, f"○ ⚠ {m}")
    else:
        _sc_item(doc, "○ ✅ 표준 구성요소 누락 없음")

    # 5. 사전 입법영향분석지표 자가평가
    _sc_h(doc, "5. 사전 입법영향분석지표 자가평가")
    if self_check.impact_assessment:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ("#", "항목", "답변", "사유")):
            _format_table_cell(cell, label, bold=True, center=True)
        for item in self_check.impact_assessment:
            row = table.add_row().cells
            ans_mark = {"yes": "● 예", "no": "● 아니오", "n/a": "— n/a"}.get(item.answer, item.answer)
            _format_table_cell(row[0], item.key, center=True)
            _format_table_cell(row[1], item.label)
            _format_table_cell(row[2], ans_mark, center=True)
            _format_table_cell(row[3], item.reason)

    # 6. 오탈자·맞춤법 점검
    _sc_h(doc, "6. 오탈자·맞춤법 점검")
    if self_check.typo_suspects:
        for t in self_check.typo_suspects:
            loc = f" ({t.location})" if t.location else ""
            sug = f" → {t.suggestion}" if t.suggestion else ""
            _sc_item(doc, f"○ ⚠ '{t.text}'{loc}{sug}")
    else:
        _sc_item(doc, "○ ✅ 의심 오탈자 없음")

    # 7. 개정안 정합성 점검 (개정안에만 표시)
    if self_check.amendment_consistency is not None:
        ac = self_check.amendment_consistency
        badge = {
            "ok": "✅ ok",
            "minor_issue": "⚠ minor_issue",
            "needs_review": "❌ needs_review",
        }.get(ac.overall, ac.overall)
        _sc_h(doc, f"7. 개정안 정합성 점검  ({badge})")
        if ac.issues:
            for iss in ac.issues:
                ref = f" ({iss.article_ref})" if iss.article_ref else ""
                _sc_item(doc, f"○ ⚠ [{iss.area}]{ref} {iss.description}")
        else:
            _sc_item(doc, "○ ✅ 지시문 ↔ 대비표 ↔ 본문 정합성 양호")


def _sc_h(doc, text: str) -> None:
    """셀프 점검 섹션 헤더 (13pt 굵게)."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _sc_item(doc, text: str) -> None:
    """○ 항목 (12pt, 들여쓰기 0.5)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(12)


def _sc_sub(doc, text: str) -> None:
    """- 소항목 (12pt, 들여쓰기 1.0)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.size = Pt(12)
