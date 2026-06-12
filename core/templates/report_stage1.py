"""Stage 1 보고서 생성기 — 마크다운 + .docx.

PolicyCard 목록을 입법 실무자(예: 경기도의회 정책지원팀)가 한눈에 의사결정할 수
있는 보고서 형태로 직렬화한다. 핵심 동선은:

  1) 경기도 기존 조례 유무 확인 → 있으면 개정, 없으면 신규 제정 검토
  2) 직접 참고(타지자체 · 일본): 원문 발췌까지 함께 본다
  3) 간접 참고(영미권): 정책 시각 참고용으로 요약만 본다

따라서 보고서 섹션 순서가 그대로 의사결정 순서를 따른다.
.docx 출력은 한글 오피스에서 그대로 열려 의회 양식과 호환된다.
"""

from __future__ import annotations

from pathlib import Path

from core.models import PolicyCard, StageReport


def _is_gyeonggi(card: PolicyCard) -> bool:
    """경기도 광역·기초 자치단체 조례인지 — jurisdiction 문자열로 단순 판별.

    kr_local 어댑터가 채우는 jurisdiction 형식은 '경기도', '경기도 ○○시', '경기도 ○○군'
    같은 표기를 따르므로 부분 문자열 매칭으로 충분하다.
    """
    return card.country == "KR" and "경기도" in (card.jurisdiction or "")


def _split_kr_cards(cards: list[PolicyCard]) -> tuple[list[PolicyCard], list[PolicyCard]]:
    """한국 조례를 (경기도, 타지자체) 두 그룹으로."""
    gg, others = [], []
    for c in cards:
        if c.country != "KR":
            continue
        if _is_gyeonggi(c):
            gg.append(c)
        else:
            others.append(c)
    return gg, others


def _render_card_with_excerpt(card: PolicyCard) -> list[str]:
    """직접 참고용 — 한국어 요약 + 원문 발췌까지 함께."""
    lines: list[str] = []
    header = f"#### {card.title}"
    if card.title_translated:
        header += f"  \n*(번역: {card.title_translated})*"
    lines.append(header)

    meta_bits: list[str] = []
    if card.jurisdiction:
        meta_bits.append(f"관할: {card.jurisdiction}")
    if card.enacted_year:
        meta_bits.append(f"제정연도: {card.enacted_year}")
    if card.url:
        meta_bits.append(f"[원문 링크]({card.url})")
    if meta_bits:
        lines.append(" · ".join(meta_bits))

    lines.append("")
    lines.append(f"**한국어 요약**: {card.summary.strip()}")

    if card.key_points:
        lines.append("")
        lines.append("**주요 포인트**")
        for kp in card.key_points:
            lines.append(f"- {kp}")

    if card.relevance_note:
        lines.append("")
        lines.append(f"**경기도 입법 시사점**: {card.relevance_note}")

    if card.raw_excerpt:
        lines.append("")
        lines.append("**원문 발췌**")
        lines.append("```")
        lines.append(card.raw_excerpt)
        lines.append("```")
    else:
        lines.append("")
        lines.append("> ⚠️ 원문을 가져오지 못했습니다 — 링크에서 직접 확인하세요.")

    lines.append("")
    return lines


def _render_card_summary_only(card: PolicyCard) -> list[str]:
    """간접 참고용 — 한국어 요약만, 원문 없이."""
    lines: list[str] = []
    header = f"#### {card.title}"
    if card.title_translated:
        header += f"  \n*(번역: {card.title_translated})*"
    lines.append(header)

    meta_bits: list[str] = []
    if card.jurisdiction:
        meta_bits.append(f"관할: {card.jurisdiction}")
    if card.enacted_year:
        meta_bits.append(f"제정연도: {card.enacted_year}")
    if card.url:
        meta_bits.append(f"[원문 링크]({card.url})")
    if meta_bits:
        lines.append(" · ".join(meta_bits))

    lines.append("")
    lines.append(f"**한국어 요약**: {card.summary.strip()}")

    if card.key_points:
        lines.append("")
        lines.append("**주요 포인트**")
        for kp in card.key_points:
            lines.append(f"- {kp}")

    if card.relevance_note:
        lines.append("")
        lines.append(f"**경기도 입법 시사점**: {card.relevance_note}")

    lines.append("")
    return lines


def _section_header(emoji: str, title: str, sub: str = "") -> list[str]:
    head = f"## {emoji} {title}"
    if sub:
        return [head, "", f"*{sub}*", ""]
    return [head, ""]


def render_stage1_markdown(report: StageReport, policy_idea: str = "") -> str:
    """Stage 1 StageReport → 마크다운 보고서 문자열."""
    if report.stage != "1":
        raise ValueError(f"Stage 1 보고서 생성기에 stage={report.stage} 리포트가 들어왔습니다.")

    cards = report.cards
    kr_gg, kr_others = _split_kr_cards(cards)
    jp_cards = [c for c in cards if c.country == "JP"]
    uk_cards = [c for c in cards if c.country == "UK"]
    eu_cards = [c for c in cards if c.country == "EU"]
    us_cards = [c for c in cards if c.country == "US"]

    out: list[str] = []
    out.append("# Stage 1 — 참고 입법례 스크리닝 보고서")
    out.append("")

    # ---------------- [개요] ----------------
    out.append("## 📋 개요")
    out.append("")
    if policy_idea:
        out.append(f"- **정책 아이디어**: {policy_idea.strip()}")
    out.append(f"- **검색 결과 요약**: {report.summary}")

    # 경기도 기존 조례 유무로 의사결정 가이드
    if kr_gg:
        out.append(
            f"- **👉 의사결정 가이드**: 경기도 기존 조례 **{len(kr_gg)}건 발견** — "
            f"먼저 본문을 검토하고 **개정안(Stage 3 amendment)** 작업이 적합한지 판단하세요."
        )
    else:
        out.append(
            "- **👉 의사결정 가이드**: 경기도 기존 조례가 **검색되지 않음** — "
            "타지자체·일본 사례를 참고해 **신규 제정안(Stage 3 new)** 작업을 검토하세요. "
            "(주의: 키워드 누락으로 검색 누락 가능성이 있으니 별도 확인 권장.)"
        )

    if report.errors:
        out.append("")
        out.append("### ⚠️ 수집 중 발생한 오류")
        for e in report.errors:
            out.append(f"- {e}")

    out.append("")

    # ---------------- [Tier A: 직접 참고] ----------------
    out.extend(
        _section_header(
            "🟢",
            "Tier A — 직접 참고 (원문 포함)",
            "조문 차용·번안이 현실적인 그룹. 본문 발췌까지 함께 확인하세요.",
        )
    )

    # 경기도
    out.append("### 🏛️ 1절. 경기도 기존 조례")
    out.append("")
    if kr_gg:
        for c in kr_gg:
            out.extend(_render_card_with_excerpt(c))
    else:
        out.append("_경기도(광역·기초) 자치법규 중 매칭되는 조례를 찾지 못했습니다._")
        out.append("")

    # 타지자체 — 카드를 먼저 순차 나열한 뒤, 절 마지막에 원문 발췌를 모아서 둠.
    # (보고서 가독성을 위해 요약·시사점만 빠르게 훑은 다음 깊이 볼 원문은 뒤로 분리.)
    out.append("### 🇰🇷 2절. 타지자체 참고 조례")
    out.append("")
    if kr_others:
        for c in kr_others:
            out.extend(_render_card_summary_only(c))
        # 카드 모두 나열 후 — 원문 발췌 모음
        out.append("#### 📜 타지자체 조례 원문 발췌 모음")
        out.append("")
        any_excerpt = False
        for idx, c in enumerate(kr_others, 1):
            if not c.raw_excerpt:
                continue
            any_excerpt = True
            out.append(f"##### 원문 {idx} — {c.title}")
            if c.url:
                out.append(f"[원문 전체 링크]({c.url})")
            out.append("")
            out.append("```")
            out.append(c.raw_excerpt)
            out.append("```")
            out.append("")
        if not any_excerpt:
            out.append("_원문을 가져오지 못한 조례들입니다 — 위 링크에서 직접 확인하세요._")
            out.append("")
    else:
        out.append("_타지자체 자치법규에서 매칭되는 조례가 없습니다._")
        out.append("")

    # 일본
    out.append("### 🇯🇵 3절. 일본 조례")
    out.append("")
    if jp_cards:
        for c in jp_cards:
            out.extend(_render_card_with_excerpt(c))
    else:
        out.append("_일본 법령에서 매칭되는 사례가 없습니다._")
        out.append("")

    # ---------------- [4절: 해외 입법례 — 간접 참고] ----------------
    out.extend(
        _section_header(
            "🟡",
            "4절. 해외 입법례 (간접 참고)",
            "법체계가 달라 직접 이식은 어렵지만 정책 시각·구조를 참고할 수 있는 그룹. 원문이 가능한 경우 링크에서 확인하세요.",
        )
    )

    for cards_block, emoji_label, fallback in [
        (uk_cards, "🇬🇧 영국 (UK)", "영국 입법례에서 매칭이 없습니다."),
        (eu_cards, "🇪🇺 EU", "EU 입법례에서 매칭이 없습니다."),
        (us_cards, "🇺🇸 미국 (US)", "미국 입법례에서 매칭이 없습니다."),
    ]:
        out.append(f"### {emoji_label}")
        out.append("")
        if cards_block:
            for c in cards_block:
                out.extend(_render_card_summary_only(c))
        else:
            out.append(f"_{fallback}_")
            out.append("")

    # ---------------- 푸터 ----------------
    out.append("---")
    out.append(
        "⚠️ 본 보고서는 자동 검색·요약 결과입니다. 인용 전 반드시 원문을 직접 확인하세요. "
        "직접 참고(Tier A) 항목에 한해 원문 발췌가 제공되며, 발췌는 길이 한도 내 일부입니다."
    )

    return "\n".join(out)


# ============ .docx renderer (공무원 표준 보고 양식) ============
#
# 양식 패턴 (경기도의회 안행위·BSC·휴가추진계획 등 분석):
#   - 1페이지: 표지(가운데 제목) + ◈ 핵심 요약 박스 + □ 개요
#   - 2페이지 이후: 절 단위로 페이지 분할, □/○/-/※ 들여쓰기 구조
#   - 마지막: 참고1, 참고2... 별도 페이지에 부속자료(조례 원문 등)
#   - 페이지 번호 "- N -" 가운데 footer
#   - 글자크기: 큰제목 14pt / 소제목 13pt / 본문 12pt (희연 팀장 표준)


def render_stage1_docx(
    report: StageReport,
    output_path: str | Path,
    policy_idea: str = "",
) -> Path:
    """Stage 1 StageReport → 공무원 표준 보고 양식의 .docx로 저장."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    KOREAN_FONT = "맑은 고딕"

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _doc_apply_default_font(doc, KOREAN_FONT, qn, OxmlElement)
    _doc_apply_margins(doc, Cm)
    _doc_add_page_number_footer(doc, OxmlElement, qn, WD_ALIGN_PARAGRAPH)

    # ---- 카드 분류 ----
    cards = report.cards
    kr_gg, kr_others = _split_kr_cards(cards)
    jp_cards = [c for c in cards if c.country == "JP"]
    uk_cards = [c for c in cards if c.country == "UK"]
    eu_cards = [c for c in cards if c.country == "EU"]
    us_cards = [c for c in cards if c.country == "US"]
    filter_notes = [e for e in report.errors if "[관련성 필터" in e]
    other_errors = [e for e in report.errors if e not in filter_notes]

    if kr_gg:
        decision_line = (
            f"경기도 기존 조례 {len(kr_gg)}건 발견 → 개정안(Stage 3 amendment) 검토 권고"
        )
    else:
        decision_line = "경기도 기존 조례 미검색 → 신규 제정안(Stage 3 new) 검토 권고"

    # ========================================
    # 1페이지: 컬러 표지 + KEY POINT 박스 + □ 개요
    # ========================================
    from datetime import date
    today_str = date.today().strftime("%Y.%m.%d")
    _doc_cover_box(
        doc,
        title="참고 입법례 스크리닝 보고서",
        subtitle=f"─ {policy_idea} ─" if policy_idea else "",
        label_no="Stage 1 Report",
        label_date=today_str,
        Pt=Pt, Cm=Cm, RGBColor=RGBColor,
        WD_ALIGN_PARAGRAPH=WD_ALIGN_PARAGRAPH, qn=qn, OxmlElement=OxmlElement,
    )

    doc.add_paragraph()

    # KEY POINT 박스
    key_point_lines = [
        f"{policy_idea or '정책 주제'}에 관한 참고 입법례 {len(cards)}건을 "
        f"{len(set(c.country for c in cards))}개 관할에서 수집·요약·분석한 결과임",
        decision_line,
    ]
    _doc_key_point_box(doc, key_point_lines, Pt, Cm, RGBColor, qn, OxmlElement)

    doc.add_paragraph()

    _doc_big(doc, "□ 개   요", Pt)
    _doc_mid(doc, f"○ 정책 주제 : {policy_idea or '(미지정)'}", Pt, Cm)
    _doc_mid(
        doc,
        f"○ 검색 결과 : 총 {len(cards)}건 ({len(set(c.country for c in cards))}개 관할)",
        Pt, Cm,
    )
    _doc_sub(
        doc,
        f"- 직접 참고 : 경기도 {len(kr_gg)}건 · 타지자체 {len(kr_others)}건 · 일본 {len(jp_cards)}건",
        Pt, Cm,
    )
    _doc_sub(
        doc,
        f"- 간접 참고 : 영국 {len(uk_cards)}건 · EU {len(eu_cards)}건 · 미국 {len(us_cards)}건",
        Pt, Cm,
    )
    _doc_mid(doc, f"○ 의사결정 : {decision_line}", Pt, Cm)
    if filter_notes:
        _doc_sub(
            doc,
            f"※ LLM 관련성 필터로 일본 {len(filter_notes)}건 자동 제외 (사유: 참고1)",
            Pt, Cm,
        )
    if not kr_gg:
        _doc_sub(doc, "※ 경기도 조례 검색 누락 가능성 별도 확인 권장", Pt, Cm)

    # ========================================
    # 2페이지: 1절. 경기도 기존 조례
    # ========================================
    _doc_pb(doc)
    _doc_section_bar(doc, "1절. 경기도 기존 조례", Pt, RGBColor, qn, OxmlElement)
    _doc_subnote(doc, "(직접 참고 — 본문 발췌 포함)", Pt)
    doc.add_paragraph()

    if kr_gg:
        _doc_big(doc, "□ 검색 결과", Pt)
        _doc_mid(doc, f"○ 매칭 조례 {len(kr_gg)}건", Pt, Cm)
        doc.add_paragraph()
        for c in kr_gg:
            _doc_render_card_full(c, doc, Pt, Cm, RGBColor, OxmlElement, qn)
    else:
        _doc_big(doc, "□ 검색 결과", Pt)
        _doc_mid(doc, "○ 경기도(광역·기초) 자치법규 중 매칭 조례 없음", Pt, Cm)
        _doc_mid(doc, "○ 권고 : 신규 제정안(Stage 3 new) 작업 검토", Pt, Cm)
        _doc_sub(doc, "※ kr_local 검색이 키워드를 놓쳤을 가능성 — 별도 확인 권장", Pt, Cm)

    # ========================================
    # 3페이지~: 2절. 타지자체 참고 조례 (요약 카드만)
    # ========================================
    _doc_pb(doc)
    _doc_section_bar(doc, "2절. 타지자체 참고 조례", Pt, RGBColor, qn, OxmlElement)
    _doc_subnote(doc, "(직접 참고 — 조례 원문은 본 절 직후 [참고1] 이하 별첨)", Pt)
    doc.add_paragraph()

    if kr_others:
        _doc_big(doc, "□ 검색 결과", Pt)
        _doc_mid(doc, f"○ 총 {len(kr_others)}건 검색", Pt, Cm)
        doc.add_paragraph()
        labels = "가나다라마바사아자차"
        for i, c in enumerate(kr_others):
            sec_label = labels[i] if i < len(labels) else str(i + 1)
            _doc_h2(doc, f"{sec_label}. {c.title}", Pt)
            _doc_render_card_meta_and_summary(c, doc, Pt, Cm, OxmlElement, qn)
            doc.add_paragraph()
    else:
        _doc_big(doc, "□ 검색 결과", Pt)
        _doc_mid(doc, "○ 타지자체 자치법규 중 매칭 조례 없음", Pt, Cm)

    # ========================================
    # 참고1~참고N: 타지자체 조례 원문 (2절 직후로 이동 — 희연 팀장 요청)
    # ========================================
    appendix_num = 1
    for c in kr_others:
        if not c.raw_excerpt:
            continue
        _doc_pb(doc)
        _doc_appendix_bar(
            doc, f"참고{appendix_num}", f"{c.title} (원문 발췌)",
            Pt, Cm, RGBColor, qn, OxmlElement,
        )
        doc.add_paragraph()
        _doc_render_appendix_excerpt(c, doc, Pt, Cm, RGBColor, OxmlElement, qn)
        appendix_num += 1

    # ========================================
    # 다음 페이지: 3절. 일본 조례
    # ========================================
    _doc_pb(doc)
    _doc_section_bar(doc, "3절. 일본 조례", Pt, RGBColor, qn, OxmlElement)
    _doc_subnote(doc, "(직접 참고 — LLM 관련성 필터 적용)", Pt)
    doc.add_paragraph()

    if jp_cards:
        _doc_big(doc, "□ 검색 결과", Pt)
        _doc_mid(doc, f"○ 관련성 필터 통과 {len(jp_cards)}건", Pt, Cm)
        doc.add_paragraph()
        labels = "가나다라마바사아자차"
        for i, c in enumerate(jp_cards):
            sec_label = labels[i] if i < len(labels) else str(i + 1)
            _doc_h2(doc, f"{sec_label}. {c.title}", Pt)
            _doc_render_card_full(c, doc, Pt, Cm, RGBColor, OxmlElement, qn)
            doc.add_paragraph()
    else:
        _doc_big(doc, "□ 검색 결과", Pt)
        _doc_mid(doc, "○ 관련성 필터링 후 매칭 없음", Pt, Cm)
        if filter_notes:
            _doc_sub(doc, "- 제외 사유 (자세히는 보고서 말미 키워드 참고):", Pt, Cm)
            for note in filter_notes:
                clean = note.replace("[관련성 필터 제외 · JP] ", "")
                _doc_sub(doc, f"  · {clean}", Pt, Cm)

    # ========================================
    # 다음 페이지: 4절. 해외 입법례
    # ========================================
    _doc_pb(doc)
    _doc_section_bar(doc, "4절. 해외 입법례", Pt, RGBColor, qn, OxmlElement)
    _doc_subnote(doc, "(간접 참고 — 정책 시각·구조 참고용, 직접 이식은 어려움)", Pt)
    doc.add_paragraph()

    _doc_big(doc, "□ 검색 결과", Pt)
    _doc_mid(
        doc,
        f"○ 영국(UK) {len(uk_cards)}건 · EU {len(eu_cards)}건 · 미국(US) {len(us_cards)}건",
        Pt, Cm,
    )
    doc.add_paragraph()

    for outer_label, code, name, block in [
        ("가", "UK", "영국 (UK)", uk_cards),
        ("나", "EU", "EU", eu_cards),
        ("다", "US", "미국 (US)", us_cards),
    ]:
        _doc_big(doc, f"□ {outer_label}. {name}", Pt)
        if not block:
            _doc_mid(doc, "○ 매칭 사례 없음", Pt, Cm)
            doc.add_paragraph()
            continue
        for c in block:
            # 카드 제목은 길 수 있어 축약
            short_title = c.title[:80] + "…" if len(c.title) > 80 else c.title
            _doc_h2(doc, f"○ {short_title}", Pt)
            _doc_render_card_meta_and_summary(c, doc, Pt, Cm, OxmlElement, qn)
            doc.add_paragraph()

    # ========================================
    # 마지막 페이지: [부록] 검색 키워드 및 필터 사유
    # ========================================
    _doc_pb(doc)
    _doc_appendix_bar(
        doc, f"참고{appendix_num}", "검색 키워드 및 관련성 필터 사유",
        Pt, Cm, RGBColor, qn, OxmlElement,
    )
    doc.add_paragraph()

    _doc_big(doc, "□ 검색 결과 요약", Pt)
    _doc_body(doc, report.summary, Pt)

    if filter_notes:
        doc.add_paragraph()
        _doc_big(doc, "□ LLM 관련성 필터 제외 사유", Pt)
        for note in filter_notes:
            clean = note.replace("[관련성 필터 제외 · JP] ", "")
            _doc_mid(doc, f"○ {clean}", Pt, Cm)

    if other_errors:
        doc.add_paragraph()
        _doc_big(doc, "□ 수집 중 발생한 오류·경고", Pt)
        for e in other_errors:
            _doc_mid(doc, f"○ {e}", Pt, Cm)

    # 푸터 안내문 — 마지막 페이지 끝에
    doc.add_paragraph()
    _doc_footer_disclaimer(
        doc,
        "⚠️ 본 보고서는 자동 검색·LLM 요약으로 생성된 결과입니다. 인용 전 반드시 원문을 "
        "직접 확인하세요. 직접 참고(Tier A) 항목에 한해 원문 발췌가 제공되며, 발췌는 길이 한도 내 일부입니다.",
        Pt,
    )

    doc.save(output_path)
    return output_path


# ---------------- docx low-level helpers ----------------


def _doc_apply_default_font(doc, font_name: str, qn, OxmlElement) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)


def _doc_apply_margins(doc, Cm) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _doc_add_page_number_footer(doc, OxmlElement, qn, WD_ALIGN_PARAGRAPH) -> None:
    """Footer에 '- N -' 형식 가운데 정렬 페이지 번호."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for child in list(p._p):
        if child.tag != qn("w:pPr"):
            p._p.remove(child)
    p.add_run("- ")
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    p.add_run(" -")


def _doc_add_hyperlink(paragraph, url: str, text: str, OxmlElement, qn, font_size: int | None = None) -> None:
    """워드 단락에 클릭 가능한 하이퍼링크 run을 추가."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if font_size is not None:
        size_el = OxmlElement("w:sz")
        size_el.set(qn("w:val"), str(font_size * 2))  # docx는 half-points 단위
        rPr.append(size_el)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _doc_pb(doc) -> None:
    """페이지 break."""
    doc.add_page_break()


def _doc_main_title(doc, text: str, Pt, WD_ALIGN_PARAGRAPH) -> None:
    """보고서 메인 제목 — 18pt 가운데 굵게."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)


def _doc_subtitle(doc, text: str, Pt, WD_ALIGN_PARAGRAPH) -> None:
    """메인 제목 아래 부제 — 13pt 가운데."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)


def _doc_h1(doc, text: str, Pt) -> None:
    """절 헤더 — 14pt 굵게 (큰제목)."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def _doc_h2(doc, text: str, Pt) -> None:
    """카드 / 하위 분류 — 13pt 굵게 (작은제목)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _doc_subnote(doc, text: str, Pt) -> None:
    """절 헤더 아래 부가 설명 — 11pt italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)


def _doc_big(doc, text: str, Pt) -> None:
    """□ 큰 항목 — 13pt 굵게, 들여쓰기 없음."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


def _doc_mid(doc, text: str, Pt, Cm) -> None:
    """○ 중 항목 — 12pt, 들여쓰기 0.5cm."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(12)


def _doc_sub(doc, text: str, Pt, Cm) -> None:
    """- 또는 ※ 소 항목 — 12pt, 들여쓰기 1.0cm."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.size = Pt(12)


def _doc_body(doc, text: str, Pt) -> None:
    """일반 본문 — 12pt."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)


def _doc_summary_box(doc, lines: list[str], Pt, Cm) -> None:
    """◈ 핵심 요약 박스 — 들여쓰기 + 굵게."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(12)


def _doc_appendix(doc, num_label: str, title: str, Pt) -> None:
    """참고N 헤더 — 14pt 굵게."""
    p = doc.add_paragraph()
    run = p.add_run(f"{num_label}    {title}")
    run.bold = True
    run.font.size = Pt(14)


def _doc_footer_disclaimer(doc, text: str, Pt) -> None:
    """문서 끝 면책 안내 — 10pt italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)


# ---- card rendering helpers ----


def _doc_render_card_meta_and_summary(
    card: PolicyCard, doc, Pt, Cm, OxmlElement, qn
) -> None:
    """카드 메타(관할·연도·링크) + 한국어 요약 + 포인트 + 시사점 (원문 발췌 제외)."""
    # 메타 라인 — 들여쓰기 0.5
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    parts: list[str] = []
    if card.jurisdiction:
        parts.append(f"관할: {card.jurisdiction}")
    if card.enacted_year:
        parts.append(f"제정연도: {card.enacted_year}")
    text_prefix = " · ".join(parts)
    if text_prefix:
        run = p.add_run(text_prefix)
        run.font.size = Pt(11)
    if card.url:
        if text_prefix:
            sep = p.add_run(" · ")
            sep.font.size = Pt(11)
        _doc_add_hyperlink(p, card.url, "원문 링크", OxmlElement, qn, font_size=11)

    if card.title_translated:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"(번역: {card.title_translated})")
        run.italic = True
        run.font.size = Pt(11)

    # ○ 한국어 요약
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    label = p.add_run("○ 한국어 요약 : ")
    label.bold = True
    label.font.size = Pt(12)
    val = p.add_run(card.summary.strip())
    val.font.size = Pt(12)

    # ○ 주요 포인트
    if card.key_points:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        label = p.add_run("○ 주요 포인트")
        label.bold = True
        label.font.size = Pt(12)
        for kp in card.key_points:
            sub_p = doc.add_paragraph()
            sub_p.paragraph_format.left_indent = Cm(1.0)
            run = sub_p.add_run(f"- {kp}")
            run.font.size = Pt(12)

    # ○ 경기도 입법 시사점
    if card.relevance_note:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        label = p.add_run("○ 경기도 입법 시사점 : ")
        label.bold = True
        label.font.size = Pt(12)
        val = p.add_run(card.relevance_note)
        val.font.size = Pt(12)


def _doc_render_card_full(
    card: PolicyCard, doc, Pt, Cm, RGBColor, OxmlElement, qn
) -> None:
    """메타·요약·포인트·시사점 + 원문 발췌까지 (1절 경기도 / 3절 일본)."""
    _doc_render_card_meta_and_summary(card, doc, Pt, Cm, OxmlElement, qn)
    if card.raw_excerpt:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        label = p.add_run("○ 원문 발췌")
        label.bold = True
        label.font.size = Pt(12)
        for line in card.raw_excerpt.split("\n"):
            if not line.strip():
                continue
            ep = doc.add_paragraph()
            ep.paragraph_format.left_indent = Cm(1.0)
            run = ep.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run("○ ※ 원문을 가져오지 못했습니다 — 위 링크에서 직접 확인하세요.")
        run.italic = True
        run.font.size = Pt(12)


def _doc_render_appendix_excerpt(
    card: PolicyCard, doc, Pt, Cm, RGBColor, OxmlElement, qn
) -> None:
    """참고 페이지에 원문 발췌 전체 표시 — 메타 한 줄 + 본문."""
    if card.jurisdiction:
        _doc_body(doc, f"○ 관할 : {card.jurisdiction}", Pt)
    if card.enacted_year:
        _doc_body(doc, f"○ 제정연도 : {card.enacted_year}", Pt)
    if card.url:
        p = doc.add_paragraph()
        label = p.add_run("○ 원문 전체 : ")
        label.font.size = Pt(12)
        _doc_add_hyperlink(p, card.url, card.url, OxmlElement, qn, font_size=12)
    doc.add_paragraph()
    if card.raw_excerpt:
        for line in card.raw_excerpt.split("\n"):
            if not line.strip():
                continue
            ep = doc.add_paragraph()
            ep.paragraph_format.left_indent = Cm(0.5)
            run = ep.add_run(line)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ============ 디자인 헬퍼 (GRI AI리포트 스타일) ============
#
# 메인 컬러: 진한 보라 (#3A2266). table cell shading, paragraph border,
# 1x1 컬러 박스로 한정된 docx 표현 안에서 디자인적 느낌을 낸다.


# 색상 상수
_BRAND_PURPLE = "3A2266"  # 메인 보라
_BRAND_PURPLE_LIGHT = "EDE7F6"  # 옅은 보라 (KEY POINT 배경)
_BRAND_GRAY_LIGHT = "F2F2F2"   # 옅은 회색 (요약 박스)


def _set_cell_bg(cell, color_hex: str, qn, OxmlElement) -> None:
    """table cell의 배경색을 16진(RRGGBB)으로 지정."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_para_left_bar(paragraph, color_hex: str, qn, OxmlElement, size_pt: int = 4) -> None:
    """단락 좌측에 컬러 세로 막대(border-left) — GRI 본문 헤더 사이드 라인 표현."""
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt * 8))  # eighths of a point
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _doc_cover_box(
    doc, title: str, subtitle: str, label_no: str, label_date: str,
    Pt, Cm, RGBColor, WD_ALIGN_PARAGRAPH, qn, OxmlElement,
) -> None:
    """표지의 컬러 제목 박스 — 보라 배경 + 흰 큰 제목 + No.라벨."""
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    cell = table.cell(0, 0)
    _set_cell_bg(cell, _BRAND_PURPLE, qn, OxmlElement)
    cell.width = Cm(16.0)

    # 본 paragraph는 cell에 이미 있음. 거기에 추가.
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 빈 줄
    p.add_run("\n")

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    if subtitle:
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p3.add_run(subtitle)
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if label_no:
        run = p4.add_run(label_no)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if label_date:
        sep = p4.add_run("   ")
        sep.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run = p4.add_run(label_date)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 박스 마지막 여백
    cell.add_paragraph().add_run(" ")


def _doc_key_point_box(
    doc, lines: list[str], Pt, Cm, RGBColor, qn, OxmlElement,
) -> None:
    """KEY POINT 라벨 박스 — 좌측 보라 라벨 + 우측 옅은 회색 본문 박스."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    label_cell = table.cell(0, 0)
    body_cell = table.cell(0, 1)
    label_cell.width = Cm(3.0)
    body_cell.width = Cm(13.0)
    _set_cell_bg(label_cell, _BRAND_PURPLE, qn, OxmlElement)
    _set_cell_bg(body_cell, _BRAND_GRAY_LIGHT, qn, OxmlElement)

    # 라벨
    p = label_cell.paragraphs[0]
    p.alignment = 1  # CENTER
    run = p.add_run("KEY\nPOINT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 본문
    first = True
    for line in lines:
        if first:
            p = body_cell.paragraphs[0]
            first = False
        else:
            p = body_cell.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _doc_section_bar(doc, text: str, Pt, RGBColor, qn, OxmlElement) -> None:
    """절 헤더 — 좌측 보라 세로 막대 + 큰 검은 제목."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    _set_para_left_bar(p, _BRAND_PURPLE, qn, OxmlElement, size_pt=6)
    run = p.add_run(f"  {text}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def _doc_appendix_bar(doc, num_label: str, title: str, Pt, Cm, RGBColor, qn, OxmlElement) -> None:
    """참고N 헤더 — 가로 컬러 박스 (좌측 라벨 + 본문)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    label_cell = table.cell(0, 0)
    body_cell = table.cell(0, 1)
    label_cell.width = Cm(2.5)
    body_cell.width = Cm(13.5)
    _set_cell_bg(label_cell, _BRAND_PURPLE, qn, OxmlElement)
    _set_cell_bg(body_cell, _BRAND_PURPLE_LIGHT, qn, OxmlElement)

    p = label_cell.paragraphs[0]
    p.alignment = 1
    run = p.add_run(num_label)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = body_cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
