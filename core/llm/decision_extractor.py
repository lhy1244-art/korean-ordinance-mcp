"""Stage 2 절충안용 — 정책 의도 + Stage 1 카드에서 '핵심 결정사항'을 추출.

사용자가 백지 명세를 쓰지 않게 하면서, 동시에 AI가 정치적/정책적 판단을
자율로 하지 않도록 하기 위한 중간 단계. LLM이 결정해야 할 항목과 선택지를
뽑고, 사용자는 선택지에만 응답하면 Stage 3 작성 시점에 그 결정값들이
반영된다.

각 선택지에는 '근거 타지자체'를 함께 적어 사용자가 판단에 활용할 수 있다 —
근거는 Stage 1 카드에 실제로 등장한 jurisdiction만 사용하도록 프롬프트에서
엄격히 강제 (환각 절대 금지).
"""

from __future__ import annotations

import json
from typing import Any

from pydantic import BaseModel, Field

from core.cache import get as cache_get, set as cache_set
from core.llm.client import get_client, model_for
from core.llm.prompts import DECISION_EXTRACTOR_SYSTEM
from core.models import PolicyCard


class DecisionOption(BaseModel):
    label: str
    text: str
    grounded_in: list[str] = Field(default_factory=list)
    note: str = ""


class DecisionCard(BaseModel):
    key: str
    question: str
    rationale: str = ""
    options: list[DecisionOption] = Field(default_factory=list)


async def extract_decision_cards(
    policy_intent: str,
    cards: list[PolicyCard],
) -> tuple[list[DecisionCard], list[str]]:
    """정책 의도 + Stage 1 카드 → 결정 카드 리스트.

    Returns:
        (decision_cards, errors). errors는 빈 경우가 일반.
    """
    if not policy_intent or not policy_intent.strip():
        return [], ["policy_intent is empty"]
    if not cards:
        return [], ["no Stage 1 cards provided"]

    # 카드 식별용 — 캐시 키 안정성 위해 url 또는 title 모음.
    card_ids = [(c.url or c.title) for c in cards]
    cache_key = {
        "task": "decision_extractor",
        "intent": policy_intent,
        "cards": card_ids,
    }
    cached = cache_get("llm", cache_key)
    if cached is not None:
        try:
            return [DecisionCard.model_validate(d) for d in cached], []
        except Exception:
            pass  # 캐시 깨진 경우 재호출

    try:
        decisions = await _ask_llm(policy_intent, cards)
    except Exception as e:  # noqa: BLE001
        return [], [f"decision_extractor: {type(e).__name__}: {e}"]

    parsed: list[DecisionCard] = []
    for d in decisions:
        if not isinstance(d, dict):
            continue
        try:
            parsed.append(DecisionCard.model_validate(d))
        except Exception:
            continue

    if parsed:
        cache_set("llm", cache_key, [d.model_dump(mode="json") for d in parsed])
    return parsed, []


async def _ask_llm(policy_intent: str, cards: list[PolicyCard]) -> list[dict[str, Any]]:
    """LLM 1회 호출 — JSON 응답에서 decisions 배열을 그대로 반환."""
    client = get_client()
    user_payload = {
        "policy_intent": policy_intent,
        "stage1_cards": [
            {
                "title": c.title,
                "country": c.country,
                "jurisdiction": c.jurisdiction,
                "summary": c.summary,
                "key_points": c.key_points,
                "raw_excerpt_head": (c.raw_excerpt or "")[:800],  # 원문 일부도 근거로 활용
            }
            for c in cards
        ],
    }
    msg = await client.messages.create(
        model=model_for("reasoning"),
        max_tokens=2500,
        system=DECISION_EXTRACTOR_SYSTEM,
        messages=[{"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)}],
    )
    text = "".join(b.text for b in msg.content if getattr(b, "type", "") == "text").strip()
    text = _strip_code_fences(text)
    data = json.loads(text)
    decisions = data.get("decisions") or []
    if not isinstance(decisions, list):
        raise ValueError("decision_extractor: 'decisions' must be a list")
    return decisions


def render_decision_cards_markdown(
    decision_cards: list[DecisionCard],
    policy_intent: str = "",
) -> str:
    """결정 카드 리스트 → 사용자에게 보여줄 마크다운 (선택지·근거 한눈에)."""
    lines: list[str] = []
    lines.append("# Stage 2 — 핵심 결정사항 카드")
    lines.append("")
    if policy_intent:
        lines.append(f"**정책 주제**: {policy_intent.strip()}")
        lines.append("")
    lines.append(
        "> 각 결정에 대해 선택지 옆 *근거 타지자체*를 확인하고 답변해 주세요. "
        "선택지 모음은 Stage 1에서 수집된 조례에서 추출된 것이므로 *환각 없음*."
    )
    lines.append("")

    for i, card in enumerate(decision_cards, 1):
        lines.append(f"## {i}. {card.question}")
        if card.rationale:
            lines.append(f"_왜 결정이 필요한가_: {card.rationale}")
            lines.append("")

        for opt in card.options:
            grounded = (
                " · ".join(opt.grounded_in) if opt.grounded_in else "_(근거 없음)_"
            )
            note_part = f" — {opt.note}" if opt.note else ""
            lines.append(f"- **{opt.label}) {opt.text}**{note_part}")
            lines.append(f"  - 근거: {grounded}")
        lines.append("")

    return "\n".join(lines)


def _strip_code_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        rows = text.splitlines()
        if rows[0].startswith("```"):
            rows = rows[1:]
        if rows and rows[-1].startswith("```"):
            rows = rows[:-1]
        text = "\n".join(rows)
    return text.strip()


# ============ docx renderer ============


def render_decision_cards_docx(
    decision_cards: list[DecisionCard],
    output_path,
    policy_intent: str = "",
):
    """결정 카드 → 워드(.docx) 파일. 공무원 양식과 폰트 통일 (14/13/12pt)."""
    from pathlib import Path

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    KOREAN_FONT = "맑은 고딕"
    BRAND_PURPLE = "3A2266"

    doc = Document()
    # 폰트
    normal = doc.styles["Normal"]
    normal.font.name = KOREAN_FONT
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    # 여백
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # 표지
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("성안 2 — 핵심 결정사항 카드")
    run.bold = True
    run.font.size = Pt(18)

    if policy_intent:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"─ {policy_intent.strip()} ─")
        run.font.size = Pt(13)

    doc.add_paragraph()

    # 안내 박스
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "EDE7F6")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    run = p.add_run(
        "◈ 각 결정에 대해 선택지 옆 근거 타지자체를 확인하고 답변하세요. "
        "선택지 모음은 1단계 스크리닝 결과에서 추출된 것이므로 환각 없음."
    )
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph()

    # 카드별
    for i, card in enumerate(decision_cards, 1):
        # 헤더
        doc.add_paragraph()
        ph = doc.add_paragraph()
        run = ph.add_run(f"□ {i}. {card.question}")
        run.bold = True
        run.font.size = Pt(14)

        # 이유
        if card.rationale:
            pr = doc.add_paragraph()
            pr.paragraph_format.left_indent = Cm(0.5)
            run = pr.add_run(f"※ 왜 결정이 필요한가 : {card.rationale}")
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # 선택지
        for opt in card.options:
            po = doc.add_paragraph()
            po.paragraph_format.left_indent = Cm(0.5)
            label_run = po.add_run(f"○ {opt.label}) {opt.text}")
            label_run.bold = True
            label_run.font.size = Pt(12)
            if opt.note:
                note_run = po.add_run(f" — {opt.note}")
                note_run.font.size = Pt(12)

            pg = doc.add_paragraph()
            pg.paragraph_format.left_indent = Cm(1.0)
            grounded = " · ".join(opt.grounded_in) if opt.grounded_in else "(근거 없음)"
            g_label = pg.add_run("근거: ")
            g_label.italic = True
            g_label.font.size = Pt(11)
            g_label.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            g_value = pg.add_run(grounded)
            g_value.font.size = Pt(11)
            g_value.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 푸터
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "⚠️ 본 결정 카드는 LLM 보조 분석 결과입니다. 정치적·정책적 판단은 사용자가 "
        "직접 결정해야 하며, 선택지에 없는 옵션도 자유롭게 추가 가능합니다."
    )
    run.italic = True
    run.font.size = Pt(10)

    doc.save(output_path)
    return output_path
