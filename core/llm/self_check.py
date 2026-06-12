"""성안 트랙의 자동 셀프 점검 (Stage 3 직후).

작성자(정책지원관 페르소나)는 *셀프 검토*를 잘 못한다는 실무 관찰에서 출발 —
초안 작성 직후 도구가 자동으로 가벼운 점검을 돌려 *주의 환기*용 결과를 첨부한다.
이건 본격 검토(별도 트랙)가 아니라, *작성자가 빠뜨린 것을 환기*하는 sanity check.

점검 항목 (프롬프트 참고):
  1. 결정값이 조문에 반영됐는지 (정합성)
  2. 상위법 명백한 저촉 (sanity)
  3. 인용 법률·조례 환각 의심
  4. 표준 구조 누락 (부칙·시행규칙 위임 등)
  5. 사전 입법영향분석지표 자가평가 (10항목 a~j)
"""

from __future__ import annotations

import json
from typing import Any, Literal

from pydantic import BaseModel, Field

from core.cache import get as cache_get, set as cache_set
from core.llm.client import get_client, model_for
from core.llm.prompts import SELF_CHECK_SYSTEM
from core.models import DraftOrdinance


class DecisionReflectionIssue(BaseModel):
    decision: str
    concern: str
    article_ref: str | None = None


class DecisionReflection(BaseModel):
    covered: list[str] = Field(default_factory=list)
    issues: list[DecisionReflectionIssue] = Field(default_factory=list)


class HigherLawConflict(BaseModel):
    article: str
    concern: str
    confidence: str = "추가검토필요"


class CitationSuspect(BaseModel):
    citation_text: str
    concern: str


class ImpactAssessmentItem(BaseModel):
    key: str
    label: str
    answer: Literal["yes", "no", "n/a"] = "n/a"
    reason: str = ""


class TypoSuspect(BaseModel):
    """오탈자·맞춤법 의심 항목."""
    text: str             # 의심 문구 그대로
    location: str = ""    # 어디에서 발견 (예: "제3조 본문")
    suggestion: str = ""  # 수정 제안 (가능하면)


class AmendmentConsistencyIssue(BaseModel):
    """개정안 정합성 점검 — 지시문 / 신·구조문대비표 / 본문 사이의 불일치."""
    area: Literal[
        "directive_vs_diff",   # 개정 지시문이 신·구조문대비표와 안 맞음
        "diff_vs_existing",    # 대비표 '현행' 칸이 기존 조례 본문과 안 맞음
        "label_mismatch",      # 조항 라벨이 본문·지시문·대비표 사이에서 다름
        "missing_in_directive",# 대비표엔 있는데 지시문엔 없음 (또는 반대)
        "addendum_issue",      # 부칙·시행일 누락 또는 모호
    ]
    description: str         # 한국어 1-2 문장
    article_ref: str = ""    # 관련 조항 (예: "제3조제4항")


class AmendmentConsistency(BaseModel):
    overall: Literal["ok", "minor_issue", "needs_review"] = "ok"
    issues: list[AmendmentConsistencyIssue] = Field(default_factory=list)


class SelfCheckResult(BaseModel):
    decision_reflection: DecisionReflection = Field(default_factory=DecisionReflection)
    higher_law_conflicts: list[HigherLawConflict] = Field(default_factory=list)
    citation_hallucination_suspects: list[CitationSuspect] = Field(default_factory=list)
    missing_standard_components: list[str] = Field(default_factory=list)
    impact_assessment: list[ImpactAssessmentItem] = Field(default_factory=list)
    typo_suspects: list[TypoSuspect] = Field(default_factory=list)
    amendment_consistency: AmendmentConsistency | None = None  # 개정안일 때만 채움
    overall_note: str = ""


async def run_self_check(
    draft: DraftOrdinance,
    decisions_text: str = "",
    policy_intent: str = "",
    *,
    mode: Literal["new", "amendment"] = "new",
    directives: list[str] | None = None,
    diffs: list[dict] | None = None,
    existing_ordinance: str = "",
) -> tuple[SelfCheckResult | None, list[str]]:
    """초안 + (사용자가 정한 결정값 텍스트) → 셀프 점검 결과.

    Args:
        draft: Stage 3 결과 DraftOrdinance.
        decisions_text: 결정 카드에서 사용자가 답변한 결정값 텍스트 (한국어 자유서술).
        policy_intent: 정책 의도 / 또는 개정 의도.
        mode: "new"(제정안) 또는 "amendment"(개정안).
        directives: 개정안 모드에서 *개정 지시문* 평문 리스트 (render에서 생성한 것 그대로).
        diffs: 개정안 모드에서 신·구조문대비표 행들 (각 행 dict 형태).
        existing_ordinance: 개정안 모드에서 기존 조례 본문 — diff_vs_existing 검증용.

    Returns:
        (result, errors). LLM 실패 시 (None, [...]). 호출자는 None이어도 파이프라인을 막지 않음.
    """
    if draft is None or not draft.articles:
        if mode != "amendment":  # 개정안은 articles 비어 있어도 changes만으로 OK
            return None, ["self_check: empty draft"]

    cache_key = {
        "task": "self_check",
        "title": draft.title,
        "mode": mode,
        "n_articles": len(draft.articles),
        "n_directives": len(directives or []),
        "n_diffs": len(diffs or []),
        "decisions": decisions_text,
    }
    cached = cache_get("llm", cache_key)
    if cached is not None:
        try:
            return SelfCheckResult.model_validate(cached), []
        except Exception:
            pass

    try:
        raw = await _ask(
            draft, decisions_text, policy_intent,
            mode=mode, directives=directives or [], diffs=diffs or [],
            existing_ordinance=existing_ordinance,
        )
    except Exception as e:  # noqa: BLE001
        return None, [f"self_check: {type(e).__name__}: {e}"]

    try:
        result = SelfCheckResult.model_validate(raw)
    except Exception as e:  # noqa: BLE001
        return None, [f"self_check parse: {type(e).__name__}: {e}"]

    cache_set("llm", cache_key, result.model_dump(mode="json"))
    return result, []


async def _ask(
    draft: DraftOrdinance,
    decisions_text: str,
    policy_intent: str,
    *,
    mode: str,
    directives: list[str],
    diffs: list[dict],
    existing_ordinance: str,
) -> dict[str, Any]:
    client = get_client()
    user_payload: dict[str, Any] = {
        "mode": mode,
        "policy_intent": policy_intent,
        "decisions_chosen": decisions_text,
        "draft": {
            "title": draft.title,
            "proposal_reason": draft.proposal_reason,
            "main_contents": draft.main_contents,
            "articles": draft.articles,
            "addendum": draft.addendum,
            "delegation_law": draft.delegation_law,
        },
    }
    if mode == "amendment":
        user_payload["amendment_directives"] = directives
        user_payload["amendment_diffs"] = diffs
        # 토큰 절약 — 기존 조례 본문 일부만 (정합성 검증에는 라벨 정도만 있어도 충분)
        user_payload["existing_ordinance_head"] = existing_ordinance[:4000]
    msg = await client.messages.create(
        model=model_for("reasoning"),
        max_tokens=4500,  # 오탈자·정합성 항목 추가로 살짝 증가
        system=SELF_CHECK_SYSTEM,
        messages=[{"role": "user", "content": json.dumps(user_payload, ensure_ascii=False)}],
    )
    text = "".join(b.text for b in msg.content if getattr(b, "type", "") == "text").strip()
    text = _strip_code_fences(text)
    return json.loads(text)


def render_self_check_markdown(result: SelfCheckResult) -> str:
    """셀프 점검 결과 → 마크다운. (필요 시 콘솔/마크다운 보고용)"""
    lines: list[str] = []
    lines.append("# 셀프 점검 결과 (참고용 — 본격 검토 아님)")
    lines.append("")
    if result.overall_note:
        lines.append(f"> **주의 환기**: {result.overall_note}")
        lines.append("")

    lines.append("## 1. 결정값 반영 점검")
    if result.decision_reflection.covered:
        lines.append("- 반영 확인: " + ", ".join(result.decision_reflection.covered))
    if result.decision_reflection.issues:
        lines.append("- ⚠️ 어긋난/누락된 결정값:")
        for issue in result.decision_reflection.issues:
            ref = f" ({issue.article_ref})" if issue.article_ref else ""
            lines.append(f"  - **{issue.decision}**{ref} — {issue.concern}")
    else:
        lines.append("- ✅ 누락·어긋남 없음")
    lines.append("")

    lines.append("## 2. 상위법 명백한 저촉 (sanity check)")
    if result.higher_law_conflicts:
        for c in result.higher_law_conflicts:
            lines.append(f"- ⚠️ {c.article}: {c.concern}  *(신뢰도: {c.confidence})*")
    else:
        lines.append("- ✅ 명백한 충돌 미확인")
    lines.append("")

    lines.append("## 3. 인용된 법률·조례 환각 의심")
    if result.citation_hallucination_suspects:
        for s in result.citation_hallucination_suspects:
            lines.append(f"- ⚠️ `{s.citation_text}` — {s.concern}")
    else:
        lines.append("- ✅ 환각 의심 인용 없음")
    lines.append("")

    lines.append("## 4. 표준 구조 누락")
    if result.missing_standard_components:
        for m in result.missing_standard_components:
            lines.append(f"- ⚠️ {m}")
    else:
        lines.append("- ✅ 표준 구성요소 누락 없음")
    lines.append("")

    lines.append("## 5. 사전 입법영향분석지표 자가평가")
    lines.append("")
    lines.append("| # | 항목 | 답변 | 사유 |")
    lines.append("|---|---|---|---|")
    for item in result.impact_assessment:
        ans = item.answer
        ans_mark = {"yes": "✅ yes", "no": "❌ no", "n/a": "— n/a"}.get(ans, ans)
        lines.append(f"| {item.key} | {item.label} | {ans_mark} | {item.reason} |")
    lines.append("")

    lines.append("## 6. 오탈자·맞춤법 점검")
    if result.typo_suspects:
        for t in result.typo_suspects:
            loc = f" ({t.location})" if t.location else ""
            sug = f" → {t.suggestion}" if t.suggestion else ""
            lines.append(f"- ⚠️ `{t.text}`{loc}{sug}")
    else:
        lines.append("- ✅ 의심 오탈자 없음")
    lines.append("")

    if result.amendment_consistency is not None:
        ac = result.amendment_consistency
        badge = {"ok": "✅ ok", "minor_issue": "⚠️ minor_issue", "needs_review": "❌ needs_review"}.get(ac.overall, ac.overall)
        lines.append(f"## 7. 개정안 정합성 점검 (overall: {badge})")
        if ac.issues:
            for iss in ac.issues:
                ref = f" ({iss.article_ref})" if iss.article_ref else ""
                lines.append(f"- ⚠️ **{iss.area}**{ref} — {iss.description}")
        else:
            lines.append("- ✅ 지시문 ↔ 대비표 ↔ 본문 정합성 양호")
        lines.append("")

    return "\n".join(lines)


def _strip_code_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        rows = text.splitlines()
        if rows[0].startswith("```"):
            rows = rows[1:]
        if rows and rows[-1].startswith("```"):
            rows = rows[:-1]
        text = "\n".join(rows)
    return text.strip()


# ============ standalone docx renderer ============


def render_self_check_docx(
    result: SelfCheckResult,
    output_path,
    draft_title: str = "",
):
    """셀프 점검 결과를 *단독* 워드(.docx) 파일로. (초안 .docx에 첨부되는 것과 별개)"""
    from pathlib import Path

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    KOREAN_FONT = "맑은 고딕"

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = KOREAN_FONT
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("성안 4 — 자동 셀프 점검 결과")
    run.bold = True
    run.font.size = Pt(18)

    if draft_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"─ 검토 대상: {draft_title} ─")
        run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("※ 본격 검토가 아닌 작성자 주의 환기용 sanity check")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # 주의 환기
    if result.overall_note:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(f"⚠ 주의 환기 : {result.overall_note}")
        run.bold = True
        run.font.size = Pt(12)

    def _h(text: str) -> None:
        doc.add_paragraph()
        ph = doc.add_paragraph()
        r = ph.add_run(text)
        r.bold = True
        r.font.size = Pt(13)

    def _item(text: str, indent: float = 0.5) -> None:
        pp = doc.add_paragraph()
        pp.paragraph_format.left_indent = Cm(indent)
        r = pp.add_run(text)
        r.font.size = Pt(12)

    # 1. 결정값 반영
    _h("1. 결정값 반영 점검")
    refl = result.decision_reflection
    if refl.covered:
        _item("○ 반영 확인: " + ", ".join(refl.covered))
    if refl.issues:
        _item("○ ⚠ 어긋남·누락된 결정값")
        for issue in refl.issues:
            ref = f" ({issue.article_ref})" if issue.article_ref else ""
            _item(f"- [{issue.decision}]{ref} {issue.concern}", indent=1.0)
    if not refl.issues:
        _item("○ ✅ 누락·어긋남 없음")

    _h("2. 상위법 명백한 저촉 (sanity check)")
    if result.higher_law_conflicts:
        for c in result.higher_law_conflicts:
            _item(f"○ ⚠ {c.article} — {c.concern}  (신뢰도: {c.confidence})")
    else:
        _item("○ ✅ 명백한 충돌 미확인")

    _h("3. 인용된 법률·조례 환각 의심")
    if result.citation_hallucination_suspects:
        for s in result.citation_hallucination_suspects:
            _item(f"○ ⚠ '{s.citation_text}' — {s.concern}")
    else:
        _item("○ ✅ 환각 의심 인용 없음")

    _h("4. 표준 구조 누락 점검")
    if result.missing_standard_components:
        for m in result.missing_standard_components:
            _item(f"○ ⚠ {m}")
    else:
        _item("○ ✅ 표준 구성요소 누락 없음")

    _h("5. 사전 입법영향분석지표 자가평가")
    if result.impact_assessment:
        table = doc.add_table(rows=1 + len(result.impact_assessment), cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for cell, label in zip(hdr, ("#", "항목", "답변", "사유")):
            ph = cell.paragraphs[0]
            ph.alignment = 1
            r = ph.add_run(label)
            r.bold = True
            r.font.size = Pt(11)
        for i, item in enumerate(result.impact_assessment, 1):
            row = table.rows[i].cells
            ans_mark = {"yes": "● 예", "no": "● 아니오", "n/a": "— n/a"}.get(item.answer, item.answer)
            for cell, text in zip(row, (item.key, item.label, ans_mark, item.reason)):
                ph = cell.paragraphs[0]
                r = ph.add_run(text)
                r.font.size = Pt(10)

    # 6. 오탈자·맞춤법
    _h("6. 오탈자·맞춤법 점검")
    if result.typo_suspects:
        for t in result.typo_suspects:
            loc = f" ({t.location})" if t.location else ""
            sug = f" → {t.suggestion}" if t.suggestion else ""
            _item(f"○ ⚠ '{t.text}'{loc}{sug}")
    else:
        _item("○ ✅ 의심 오탈자 없음")

    # 7. 개정안 정합성 (개정안일 때만)
    if result.amendment_consistency is not None:
        ac = result.amendment_consistency
        badge = {"ok": "✅ ok", "minor_issue": "⚠ minor_issue", "needs_review": "❌ needs_review"}.get(ac.overall, ac.overall)
        _h(f"7. 개정안 정합성 점검  ({badge})")
        if ac.issues:
            for iss in ac.issues:
                ref = f" ({iss.article_ref})" if iss.article_ref else ""
                _item(f"○ ⚠ [{iss.area}]{ref} {iss.description}")
        else:
            _item("○ ✅ 지시문 ↔ 대비표 ↔ 본문 정합성 양호")

    # 푸터
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "⚠️ 본 셀프 점검은 LLM 보조 분석입니다. 본격 검토(상위법·타지자체 비교 등)는 "
        "검토 트랙에서 수행하며, 입법조사관 최종 확인이 필요합니다."
    )
    run.italic = True
    run.font.size = Pt(10)

    doc.save(output_path)
    return output_path
